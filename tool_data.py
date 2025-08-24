import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import seaborn as sns
from pathlib import Path
import json
import io
import base64
from typing import Dict, List, Any, Optional
from docx import Document
import openpyxl
from datetime import datetime
import re

# Set matplotlib backend for headless environments
plt.switch_backend('Agg')
sns.set_style("whitegrid")

class DataProcessor:
    """Enhanced data processor with file upload and visualization capabilities"""
    
    def __init__(self):
        self.supported_formats = ['.xlsx', '.xls', '.csv', '.docx', '.txt', '.json']
        self.upload_dir = Path("uploads")
        self.upload_dir.mkdir(exist_ok=True)
    
    def process_file(self, file_path: str) -> Dict[str, Any]:
        """Process uploaded file and extract data"""
        file_path = Path(file_path)
        
        if not file_path.exists():
            return {"error": "File not found", "summary": "File processing failed"}
        
        try:
            if file_path.suffix.lower() in ['.xlsx', '.xls']:
                return self._process_excel(file_path)
            elif file_path.suffix.lower() == '.csv':
                return self._process_csv(file_path)
            elif file_path.suffix.lower() == '.docx':
                return self._process_docx(file_path)
            elif file_path.suffix.lower() == '.json':
                return self._process_json(file_path)
            elif file_path.suffix.lower() == '.txt':
                return self._process_text(file_path)
            else:
                return {"error": f"Unsupported format: {file_path.suffix}", "summary": "Format not supported"}
        
        except Exception as e:
            return {"error": str(e), "summary": f"Error processing file: {str(e)}"}
    
    def _process_excel(self, file_path: Path) -> Dict[str, Any]:
        """Process Excel files with multiple sheets"""
        try:
            excel_file = pd.ExcelFile(file_path)
            sheets_data = {}
            all_data = []
            
            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(file_path, sheet_name=sheet_name)
                sheets_data[sheet_name] = df
                all_data.append(df)
            
            if all_data:
                combined_df = pd.concat(all_data, ignore_index=True)
                return self._analyze_dataframe(combined_df, sheets_data, "Excel")
            
        except Exception as e:
            return {"error": str(e), "summary": "Failed to process Excel file"}
    
    def _process_csv(self, file_path: Path) -> Dict[str, Any]:
        """Process CSV files"""
        try:
            df = pd.read_csv(file_path)
            return self._analyze_dataframe(df, {"main": df}, "CSV")
        except Exception as e:
            return {"error": str(e), "summary": "Failed to process CSV file"}
    
    def _process_docx(self, file_path: Path) -> Dict[str, Any]:
        """Extract tables and text from Word documents"""
        try:
            doc = Document(file_path)
            
            # Extract text
            text_content = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text.strip())
            
            # Extract tables
            tables_data = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(row_data)
                
                if table_data:
                    try:
                        df = pd.DataFrame(table_data[1:], columns=table_data[0])
                        tables_data.append(df)
                    except:
                        tables_data.append(table_data)
            
            # Analyze extracted data
            if tables_data and isinstance(tables_data[0], pd.DataFrame):
                combined_df = pd.concat([df for df in tables_data if isinstance(df, pd.DataFrame)], ignore_index=True)
                result = self._analyze_dataframe(combined_df, {"tables": tables_data}, "Word Document")
                result["text_content"] = text_content[:10]
                return result
            else:
                return {
                    "summary": f"Extracted {len(text_content)} paragraphs and {len(tables_data)} tables from Word document",
                    "text_content": text_content[:10],
                    "tables_count": len(tables_data),
                    "raw": {"text": text_content, "tables": tables_data}
                }
                
        except Exception as e:
            return {"error": str(e), "summary": "Failed to process Word document"}
    
    def _process_json(self, file_path: Path) -> Dict[str, Any]:
        """Process JSON files"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
            
            if isinstance(data, list) and data and isinstance(data[0], dict):
                df = pd.DataFrame(data)
                return self._analyze_dataframe(df, {"main": df}, "JSON")
            else:
                return {
                    "summary": f"JSON file contains {type(data).__name__} with {len(data) if hasattr(data, '__len__') else 'unknown'} items",
                    "raw": data,
                    "data_type": type(data).__name__
                }
                
        except Exception as e:
            return {"error": str(e), "summary": "Failed to process JSON file"}
    
    def _process_text(self, file_path: Path) -> Dict[str, Any]:
        """Process text files and try to extract structured data"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            lines = content.split('\n')
            non_empty_lines = [line.strip() for line in lines if line.strip()]
            
            # Try to detect CSV-like structure
            if any(',' in line or '\t' in line for line in non_empty_lines[:5]):
                try:
                    df = pd.read_csv(io.StringIO(content))
                    return self._analyze_dataframe(df, {"main": df}, "Text (CSV-like)")
                except:
                    pass
            
            # Extract numbers for basic analysis
            numbers = []
            for line in non_empty_lines:
                found_numbers = re.findall(r'-?\d+\.?\d*', line)
                numbers.extend([float(n) for n in found_numbers if n])
            
            summary = f"Text file with {len(non_empty_lines)} lines"
            if numbers:
                summary += f", extracted {len(numbers)} numbers (avg: {np.mean(numbers):.2f})"
            
            return {
                "summary": summary,
                "lines_count": len(non_empty_lines),
                "numbers_found": len(numbers),
                "numbers_stats": {
                    "count": len(numbers),
                    "mean": float(np.mean(numbers)) if numbers else 0,
                    "sum": float(np.sum(numbers)) if numbers else 0
                } if numbers else None,
                "raw": {"content": content[:1000], "numbers": numbers[:50]}
            }
            
        except Exception as e:
            return {"error": str(e), "summary": "Failed to process text file"}
    
    def _analyze_dataframe(self, df: pd.DataFrame, sheets_data: Dict, source_type: str) -> Dict[str, Any]:
        """Comprehensive analysis of DataFrame with visualizations"""
        try:
            # Clean DataFrame for JSON serialization
            df_clean = df.fillna("")  # Replace NaN with empty strings
            
            analysis = {
                "source_type": source_type,
                "shape": df.shape,
                "columns": list(df.columns),
                "summary": "",
                "statistics": {},
                "visualizations": [],
                "insights": [],
                "raw": df_clean.head(100).to_dict('records')
            }
            
            # Basic statistics
            numeric_cols = df.select_dtypes(include=[np.number]).columns.tolist()
            categorical_cols = df.select_dtypes(include=['object']).columns.tolist()
            
            # Convert missing values to regular integers for JSON serialization
            missing_values = df.isnull().sum().to_dict()
            missing_values_clean = {k: int(v) for k, v in missing_values.items()}
            
            analysis["statistics"] = {
                "numeric_columns": len(numeric_cols),
                "categorical_columns": len(categorical_cols),
                "total_rows": len(df),
                "missing_values": missing_values_clean
            }
            
            # Generate insights
            insights = []
            insights.append(f"Dataset contains {len(df)} rows and {len(df.columns)} columns")
            
            if numeric_cols:
                insights.append(f"Found {len(numeric_cols)} numeric columns: {', '.join(numeric_cols[:3])}{'...' if len(numeric_cols) > 3 else ''}")
                
                # Clean numeric statistics for JSON serialization
                numeric_stats_raw = df[numeric_cols].describe().to_dict()
                numeric_stats_clean = {}
                for col, stats in numeric_stats_raw.items():
                    numeric_stats_clean[col] = {}
                    for stat_name, stat_value in stats.items():
                        try:
                            if pd.isna(stat_value):
                                numeric_stats_clean[col][stat_name] = None
                            elif isinstance(stat_value, (int, float, np.integer, np.floating)) and np.isinf(stat_value):
                                numeric_stats_clean[col][stat_name] = None
                            else:
                                numeric_stats_clean[col][stat_name] = float(stat_value)
                        except (TypeError, ValueError):
                            numeric_stats_clean[col][stat_name] = str(stat_value)
                
                analysis["numeric_statistics"] = numeric_stats_clean
                
                for col in numeric_cols[:3]:
                    col_data = df[col].dropna()
                    if len(col_data) > 0:
                        try:
                            mean_val = col_data.mean()
                            min_val = col_data.min()
                            max_val = col_data.max()
                            
                            # Check for valid numeric values
                            if not (pd.isna(mean_val) or (isinstance(mean_val, (int, float, np.integer, np.floating)) and np.isinf(mean_val))):
                                insights.append(f"{col}: mean={mean_val:.2f}, range={min_val:.2f}-{max_val:.2f}")
                        except (TypeError, ValueError) as e:
                            insights.append(f"{col}: contains mixed or non-numeric data")
            
            if categorical_cols:
                insights.append(f"Found {len(categorical_cols)} categorical columns")
                
                for col in categorical_cols[:2]:
                    unique_vals = df[col].nunique()
                    insights.append(f"{col}: {unique_vals} unique values")
            
            analysis["insights"] = insights
            
            # Generate visualizations
            visualizations = self._create_visualizations(df, numeric_cols, categorical_cols)
            analysis["visualizations"] = visualizations
            
            # Create summary
            summary_parts = [
                f"Processed {source_type} file with {len(df)} records",
                f"Contains {len(numeric_cols)} numeric and {len(categorical_cols)} text columns"
            ]
            
            if numeric_cols:
                try:
                    total_sum = df[numeric_cols].sum().sum()
                    if not (pd.isna(total_sum) or (isinstance(total_sum, (int, float, np.integer, np.floating)) and np.isinf(total_sum))):
                        summary_parts.append(f"Total numeric sum: {total_sum:,.2f}")
                except (TypeError, ValueError):
                    summary_parts.append("Contains mixed numeric data types")
            
            analysis["summary"] = ". ".join(summary_parts)
            
            return analysis
            
        except Exception as e:
            return {"error": str(e), "summary": f"Failed to analyze {source_type} data"}
    
    def _create_visualizations(self, df: pd.DataFrame, numeric_cols: List[str], categorical_cols: List[str]) -> List[Dict[str, str]]:
        """Create visualizations and return as base64 encoded images"""
        visualizations = []
        
        try:
            # 1. Numeric columns distribution
            if numeric_cols:
                fig, axes = plt.subplots(min(2, len(numeric_cols)), 1, figsize=(10, 6))
                if len(numeric_cols) == 1:
                    axes = [axes]
                
                for i, col in enumerate(numeric_cols[:2]):
                    ax = axes[i] if len(numeric_cols) > 1 else axes[0]
                    df[col].hist(bins=20, ax=ax, alpha=0.7, color='skyblue')
                    ax.set_title(f'Distribution of {col}')
                    ax.set_xlabel(col)
                    ax.set_ylabel('Frequency')
                
                plt.tight_layout()
                img_buffer = io.BytesIO()
                plt.savefig(img_buffer, format='png', dpi=100, bbox_inches='tight')
                img_buffer.seek(0)
                img_b64 = base64.b64encode(img_buffer.getvalue()).decode()
                plt.close()
                
                visualizations.append({
                    "title": "Numeric Distributions",
                    "type": "histogram",
                    "image": img_b64
                })
            
            # 2. Correlation heatmap (if multiple numeric columns)
            if len(numeric_cols) > 1:
                plt.figure(figsize=(8, 6))
                correlation_matrix = df[numeric_cols].corr()
                sns.heatmap(correlation_matrix, annot=True, cmap='coolwarm', center=0)
                plt.title('Correlation Matrix')
                plt.tight_layout()
                
                img_buffer = io.BytesIO()
                plt.savefig(img_buffer, format='png', dpi=100, bbox_inches='tight')
                img_buffer.seek(0)
                img_b64 = base64.b64encode(img_buffer.getvalue()).decode()
                plt.close()
                
                visualizations.append({
                    "title": "Correlation Matrix",
                    "type": "heatmap",
                    "image": img_b64
                })
            
            # 3. Top categories (if categorical columns exist)
            if categorical_cols:
                col = categorical_cols[0]
                top_categories = df[col].value_counts().head(10)
                
                plt.figure(figsize=(10, 6))
                top_categories.plot(kind='bar', color='lightcoral')
                plt.title(f'Top 10 {col} Categories')
                plt.xlabel(col)
                plt.ylabel('Count')
                plt.xticks(rotation=45)
                plt.tight_layout()
                
                img_buffer = io.BytesIO()
                plt.savefig(img_buffer, format='png', dpi=100, bbox_inches='tight')
                img_buffer.seek(0)
                img_b64 = base64.b64encode(img_buffer.getvalue()).decode()
                plt.close()
                
                visualizations.append({
                    "title": f"Top {col} Categories",
                    "type": "bar_chart",
                    "image": img_b64
                })
        
        except Exception as e:
            print(f"Error creating visualizations: {e}")
        
        return visualizations

# Helper function to clean data for JSON serialization
def clean_for_json(obj):
    """Recursively clean data structures to be JSON serializable"""
    if isinstance(obj, dict):
        return {k: clean_for_json(v) for k, v in obj.items()}
    elif isinstance(obj, list):
        return [clean_for_json(item) for item in obj]
    elif pd.isna(obj):
        return None
    elif isinstance(obj, (int, float, np.integer, np.floating)):
        try:
            if np.isnan(obj) or np.isinf(obj):
                return None
            elif isinstance(obj, (np.integer, np.int64)):
                return int(obj)
            elif isinstance(obj, (np.floating, np.float64)):
                return float(obj)
            else:
                return obj
        except (TypeError, ValueError):
            return str(obj)  # Convert problematic numeric types to string
    else:
        return obj

# Global processor instance
data_processor = DataProcessor()

def fetch_and_summarize_data(source=None, file_path=None):
    """
    Enhanced data fetching with file upload support
    
    Args:
        source: Optional data source identifier
        file_path: Path to uploaded file for processing
    
    Returns:
        Dict with summary, visualizations, and raw data
    """
    
    if file_path:
        # Process uploaded file
        result = data_processor.process_file(file_path)
        result["timestamp"] = datetime.now().isoformat()
        result["source"] = f"file:{Path(file_path).name}"
        # Clean the result for JSON serialization
        return clean_for_json(result)
    
    # Fallback to demo data if no file provided
    demo_data = [
        {"date": "2025-01-19", "sales": 1200, "region": "North", "product": "Widget A"},
        {"date": "2025-01-20", "sales": 1450, "region": "South", "product": "Widget B"},
        {"date": "2025-01-21", "sales": 980, "region": "East", "product": "Widget A"},
        {"date": "2025-01-22", "sales": 1650, "region": "West", "product": "Widget C"},
        {"date": "2025-01-23", "sales": 1320, "region": "North", "product": "Widget B"}
    ]
    
    df = pd.DataFrame(demo_data)
    total = df["sales"].sum()
    avg_sales = df["sales"].mean()
    
    return {
        "summary": f"Demo data: Total sales of {total:,} across {len(demo_data)} days (avg: {avg_sales:.0f} per day)",
        "statistics": {
            "total_sales": total,
            "average_sales": avg_sales,
            "days": len(demo_data),
            "regions": df["region"].nunique(),
            "products": df["product"].nunique()
        },
        "raw": demo_data,
        "source": "demo_data",
        "timestamp": datetime.now().isoformat()
    }
